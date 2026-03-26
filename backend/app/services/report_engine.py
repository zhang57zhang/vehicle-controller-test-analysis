"""
报告生成引擎

提供测试报告生成功能：
- PDF报告生成
- Word报告生成
- 报告模板管理
- 数据溯源信息记录
"""

import logging
import os
import json
import io
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, black, white, grey
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Image,
    ListFlowable,
    ListItem,
)
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.config import settings

logger = logging.getLogger(__name__)


class ReportGenerationError(Exception):
    """报告生成错误"""

    pass


class ReportEngine:
    """
    报告生成引擎

    支持生成标准报告和溯源报告
    """

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """设置自定义样式"""
        self.styles.add(
            ParagraphStyle(
                name="ChineseTitle",
                parent=self.styles["Title"],
                fontSize=24,
                leading=30,
                alignment=TA_CENTER,
                spaceAfter=20,
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseHeading1",
                parent=self.styles["Heading1"],
                fontSize=16,
                leading=20,
                spaceBefore=15,
                spaceAfter=10,
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseHeading2",
                parent=self.styles["Heading2"],
                fontSize=14,
                leading=18,
                spaceBefore=12,
                spaceAfter=8,
            )
        )

        self.styles.add(
            ParagraphStyle(
                name="ChineseBody",
                parent=self.styles["Normal"],
                fontSize=10,
                leading=14,
                spaceBefore=6,
                spaceAfter=6,
            )
        )

    def generate_standard_report(
        self, report_data: Dict[str, Any], output_path: str, format: str = "pdf"
    ) -> str:
        """
        生成标准报告（对外）

        Args:
            report_data: 报告数据
            output_path: 输出路径
            format: 输出格式 ('pdf' 或 'word')

        Returns:
            生成的报告文件路径
        """
        if format == "pdf":
            return self._generate_pdf_standard_report(report_data, output_path)
        else:
            return self._generate_word_standard_report(report_data, output_path)

    def generate_traceability_report(
        self, report_data: Dict[str, Any], output_path: str, format: str = "pdf"
    ) -> str:
        """
        生成溯源报告（内部审核）

        Args:
            report_data: 报告数据
            output_path: 输出路径
            format: 输出格式 ('pdf' 或 'word')

        Returns:
            生成的报告文件路径
        """
        if format == "pdf":
            return self._generate_pdf_traceability_report(report_data, output_path)
        else:
            return self._generate_word_traceability_report(report_data, output_path)

    def _generate_pdf_standard_report(self, report_data: Dict[str, Any], output_path: str) -> str:
        """生成PDF标准报告"""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        story = []

        story.extend(self._create_cover_page(report_data))
        story.append(PageBreak())

        story.extend(self._create_toc(report_data))
        story.append(PageBreak())

        story.extend(self._create_summary_section(report_data))

        story.extend(self._create_results_section(report_data))

        story.extend(self._create_indicators_section(report_data))

        story.extend(self._create_issues_section(report_data))

        doc.build(story, onFirstPage=self._add_page_number)

        return output_path

    def _create_cover_page(self, report_data: Dict[str, Any]) -> List:
        """创建封面页"""
        elements = []

        elements.append(Spacer(1, 3 * cm))

        title = report_data.get("project_name", "测试报告")
        elements.append(Paragraph(title, self.styles["ChineseTitle"]))

        elements.append(Spacer(1, 2 * cm))

        cover_info = [
            ["报告编号", report_data.get("report_number", "N/A")],
            ["报告日期", report_data.get("report_date", datetime.now().strftime("%Y-%m-%d"))],
            ["版本", report_data.get("version", "v1.0")],
            ["编制", report_data.get("author", "")],
            ["审核", report_data.get("reviewer", "")],
            ["批准", report_data.get("approver", "")],
        ]

        table = Table(cover_info, colWidths=[4 * cm, 8 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 11),
                    ("ALIGN", (0, 0), (0, -1), "RIGHT"),
                    ("ALIGN", (1, 0), (1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )

        elements.append(table)

        return elements

    def _create_toc(self, report_data: Dict[str, Any]) -> List:
        """创建目录页"""
        elements = []

        elements.append(Paragraph("目录", self.styles["ChineseHeading1"]))
        elements.append(Spacer(1, 0.5 * cm))

        toc_items = ["1. 测试概述", "2. 测试结果汇总", "3. 关键指标分析", "4. 问题列表", "5. 附录"]

        for item in toc_items:
            elements.append(Paragraph(item, self.styles["ChineseBody"]))

        return elements

    def _create_summary_section(self, report_data: Dict[str, Any]) -> List:
        """创建测试概述章节"""
        elements = []

        elements.append(Paragraph("1. 测试概述", self.styles["ChineseHeading1"]))

        elements.append(Paragraph("1.1 测试目的", self.styles["ChineseHeading2"]))
        purpose = report_data.get("test_purpose", "验证控制器功能是否符合设计要求。")
        elements.append(Paragraph(purpose, self.styles["ChineseBody"]))

        elements.append(Paragraph("1.2 测试范围", self.styles["ChineseHeading2"]))
        scope = report_data.get("test_scope", "功能测试、性能测试")
        elements.append(Paragraph(scope, self.styles["ChineseBody"]))

        elements.append(Paragraph("1.3 测试环境", self.styles["ChineseHeading2"]))

        env_info = [
            ["测试类型", report_data.get("test_type", "HIL")],
            ["测试日期", report_data.get("test_date", datetime.now().strftime("%Y-%m-%d"))],
            ["软件版本", report_data.get("software_version", "N/A")],
            ["硬件版本", report_data.get("hardware_version", "N/A")],
            ["DBC版本", report_data.get("dbc_version", "N/A")],
        ]

        table = Table(env_info, colWidths=[4 * cm, 10 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )

        elements.append(table)
        elements.append(Spacer(1, 0.5 * cm))

        return elements

    def _create_results_section(self, report_data: Dict[str, Any]) -> List:
        """创建测试结果汇总章节"""
        elements = []

        elements.append(Paragraph("2. 测试结果汇总", self.styles["ChineseHeading1"]))

        summary = report_data.get("summary", {})
        total = summary.get("total", 0)
        passed = summary.get("pass", 0)
        failed = summary.get("fail", 0)
        warning = summary.get("warning", 0)

        pass_rate = (passed / total * 100) if total > 0 else 0

        result_text = f"本次测试共执行 {total} 项测试用例，其中通过 {passed} 项，失败 {failed} 项，警告 {warning} 项。通过率为 {pass_rate:.1f}%。"
        elements.append(Paragraph(result_text, self.styles["ChineseBody"]))

        elements.append(Spacer(1, 0.3 * cm))

        summary_table = [
            ["指标", "数量", "占比"],
            ["总用例数", str(total), "100%"],
            ["通过", str(passed), f"{pass_rate:.1f}%"],
            ["失败", str(failed), f"{(failed / total * 100) if total > 0 else 0:.1f}%"],
            ["警告", str(warning), f"{(warning / total * 100) if total > 0 else 0:.1f}%"],
        ]

        table = Table(summary_table, colWidths=[5 * cm, 4 * cm, 4 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1890ff")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )

        elements.append(table)
        elements.append(Spacer(1, 0.5 * cm))

        return elements

    def _create_indicators_section(self, report_data: Dict[str, Any]) -> List:
        """创建关键指标分析章节"""
        elements = []

        elements.append(Paragraph("3. 关键指标分析", self.styles["ChineseHeading1"]))

        indicators = report_data.get("indicators", [])

        if not indicators:
            elements.append(Paragraph("暂无指标数据", self.styles["ChineseBody"]))
            return elements

        indicator_table = [["指标ID", "类型", "信号", "状态", "备注"]]

        for ind in indicators[:20]:
            indicator_table.append(
                [
                    str(ind.get("indicator_id", ""))[:15],
                    str(ind.get("indicator_type", ""))[:10],
                    str(ind.get("signal_name", ""))[:20],
                    str(ind.get("result_status", "")),
                    str(ind.get("notes", ""))[:30],
                ]
            )

        table = Table(indicator_table, colWidths=[2.5 * cm, 2 * cm, 4 * cm, 2 * cm, 4 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1890ff")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        elements.append(table)
        elements.append(Spacer(1, 0.5 * cm))

        return elements

    def _create_issues_section(self, report_data: Dict[str, Any]) -> List:
        """创建问题列表章节"""
        elements = []

        elements.append(Paragraph("4. 问题列表", self.styles["ChineseHeading1"]))

        issues = report_data.get("issues", [])

        if not issues:
            elements.append(Paragraph("本次测试未发现问题。", self.styles["ChineseBody"]))
            return elements

        issue_table = [["序号", "问题描述", "严重程度", "状态"]]

        for i, issue in enumerate(issues, 1):
            issue_table.append(
                [
                    str(i),
                    str(issue.get("description", ""))[:40],
                    str(issue.get("severity", "中")),
                    str(issue.get("status", "待处理")),
                ]
            )

        table = Table(issue_table, colWidths=[1.5 * cm, 8 * cm, 2.5 * cm, 2.5 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ff4d4f")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        elements.append(table)

        return elements

    def _add_page_number(self, canvas, doc):
        """添加页码"""
        page_num = canvas.getPageNumber()
        text = f"第 {page_num} 页"
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(A4[0] / 2, 1 * cm, text)
        canvas.restoreState()

    def _generate_pdf_traceability_report(
        self, report_data: Dict[str, Any], output_path: str
    ) -> str:
        """生成PDF溯源报告"""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        story = []

        story.extend(self._create_cover_page(report_data))
        story.append(PageBreak())

        story.extend(self._create_data_source_section(report_data))

        story.extend(self._create_signal_traceability_section(report_data))

        story.extend(self._create_detailed_results_section(report_data))

        story.extend(self._create_raw_data_reference_section(report_data))

        doc.build(story, onFirstPage=self._add_page_number)

        return output_path

    def _create_data_source_section(self, report_data: Dict[str, Any]) -> List:
        """创建数据来源章节"""
        elements = []

        elements.append(Paragraph("1. 数据来源信息", self.styles["ChineseHeading1"]))

        data_sources = report_data.get("data_sources", [])

        source_table = [["文件名", "文件类型", "大小", "导入时间", "数据质量"]]

        for source in data_sources:
            source_table.append(
                [
                    str(source.get("file_name", ""))[:25],
                    str(source.get("file_type", "")),
                    str(source.get("file_size", "")),
                    str(source.get("import_time", "")),
                    str(source.get("data_quality", "valid")),
                ]
            )

        table = Table(source_table, colWidths=[5 * cm, 2 * cm, 2 * cm, 3 * cm, 2 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#52c41a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        elements.append(table)
        elements.append(Spacer(1, 0.5 * cm))

        return elements

    def _create_signal_traceability_section(self, report_data: Dict[str, Any]) -> List:
        """创建信号溯源章节"""
        elements = []

        elements.append(Paragraph("2. 信号溯源表", self.styles["ChineseHeading1"]))

        signal_mappings = report_data.get("signal_mappings", [])

        mapping_table = [["信号别名", "原始信号", "DBC信号", "单位转换", "数据来源"]]

        for mapping in signal_mappings:
            unit_conv = mapping.get("unit_conversion", {})
            conv_str = ""
            if unit_conv:
                conv_str = f"{unit_conv.get('from_unit', '')} -> {unit_conv.get('to_unit', '')}"

            mapping_table.append(
                [
                    str(mapping.get("signal_alias", ""))[:15],
                    str(mapping.get("data_source_signal", ""))[:15],
                    str(mapping.get("dbc_signal", ""))[:15],
                    conv_str[:15],
                    str(mapping.get("source_file", ""))[:15],
                ]
            )

        table = Table(mapping_table, colWidths=[3 * cm, 3 * cm, 3 * cm, 3 * cm, 3 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#722ed1")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )

        elements.append(table)
        elements.append(Spacer(1, 0.5 * cm))

        return elements

    def _create_detailed_results_section(self, report_data: Dict[str, Any]) -> List:
        """创建详细结果章节"""
        elements = []

        elements.append(Paragraph("3. 详细分析结果", self.styles["ChineseHeading1"]))

        indicators = report_data.get("indicators", [])

        for ind in indicators[:10]:
            elements.append(
                Paragraph(
                    f"3.{indicators.index(ind) + 1} {ind.get('indicator_id', 'Unknown')}",
                    self.styles["ChineseHeading2"],
                )
            )

            elements.append(
                Paragraph(f"类型: {ind.get('indicator_type', 'N/A')}", self.styles["ChineseBody"])
            )
            elements.append(
                Paragraph(f"状态: {ind.get('result_status', 'N/A')}", self.styles["ChineseBody"])
            )

            result_value = ind.get("result_value", {})
            if isinstance(result_value, dict):
                for key, value in result_value.items():
                    elements.append(Paragraph(f"  - {key}: {value}", self.styles["ChineseBody"]))

            elements.append(Paragraph(f"备注: {ind.get('notes', '')}", self.styles["ChineseBody"]))
            elements.append(Spacer(1, 0.3 * cm))

        return elements

    def _create_raw_data_reference_section(self, report_data: Dict[str, Any]) -> List:
        """创建原始数据引用章节"""
        elements = []

        elements.append(Paragraph("4. 原始数据引用", self.styles["ChineseHeading1"]))

        elements.append(
            Paragraph(
                "本章节记录所有分析结果对应的原始数据位置，便于追溯和验证。",
                self.styles["ChineseBody"],
            )
        )

        provenance = report_data.get("provenance", {})

        for key, value in provenance.items():
            elements.append(Paragraph(f"  - {key}: {value}", self.styles["ChineseBody"]))

        return elements

    def _generate_word_standard_report(self, report_data: Dict[str, Any], output_path: str) -> str:
        """生成Word标准报告"""
        doc = Document()

        title = doc.add_heading(report_data.get("project_name", "测试报告"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        cover_table = doc.add_table(rows=6, cols=2)
        cover_table.style = "Table Grid"

        cover_data = [
            ("报告编号", report_data.get("report_number", "N/A")),
            ("报告日期", report_data.get("report_date", datetime.now().strftime("%Y-%m-%d"))),
            ("版本", report_data.get("version", "v1.0")),
            ("编制", report_data.get("author", "")),
            ("审核", report_data.get("reviewer", "")),
            ("批准", report_data.get("approver", "")),
        ]

        for i, (label, value) in enumerate(cover_data):
            row = cover_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

        doc.add_page_break()

        doc.add_heading("1. 测试概述", level=1)

        doc.add_heading("1.1 测试目的", level=2)
        doc.add_paragraph(report_data.get("test_purpose", "验证控制器功能是否符合设计要求。"))

        doc.add_heading("1.2 测试环境", level=2)

        env_table = doc.add_table(rows=5, cols=2)
        env_table.style = "Table Grid"

        env_data = [
            ("测试类型", report_data.get("test_type", "HIL")),
            ("测试日期", report_data.get("test_date", datetime.now().strftime("%Y-%m-%d"))),
            ("软件版本", report_data.get("software_version", "N/A")),
            ("硬件版本", report_data.get("hardware_version", "N/A")),
            ("DBC版本", report_data.get("dbc_version", "N/A")),
        ]

        for i, (label, value) in enumerate(env_data):
            row = env_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

        doc.add_heading("2. 测试结果汇总", level=1)

        summary = report_data.get("summary", {})
        total = summary.get("total", 0)
        passed = summary.get("pass", 0)
        failed = summary.get("fail", 0)
        warning = summary.get("warning", 0)
        pass_rate = (passed / total * 100) if total > 0 else 0

        doc.add_paragraph(
            f"本次测试共执行 {total} 项测试用例，其中通过 {passed} 项，"
            f"失败 {failed} 项，警告 {warning} 项。通过率为 {pass_rate:.1f}%。"
        )

        doc.add_heading("3. 关键指标分析", level=1)

        indicators = report_data.get("indicators", [])
        if indicators:
            ind_table = doc.add_table(rows=min(len(indicators) + 1, 21), cols=5)
            ind_table.style = "Table Grid"

            header = ind_table.rows[0]
            header.cells[0].text = "指标ID"
            header.cells[1].text = "类型"
            header.cells[2].text = "信号"
            header.cells[3].text = "状态"
            header.cells[4].text = "备注"

            for i, ind in enumerate(indicators[:20]):
                row = ind_table.rows[i + 1]
                row.cells[0].text = str(ind.get("indicator_id", ""))[:15]
                row.cells[1].text = str(ind.get("indicator_type", ""))[:10]
                row.cells[2].text = str(ind.get("signal_name", ""))[:20]
                row.cells[3].text = str(ind.get("result_status", ""))
                row.cells[4].text = str(ind.get("notes", ""))[:30]

        doc.add_heading("4. 问题列表", level=1)

        issues = report_data.get("issues", [])
        if issues:
            issue_table = doc.add_table(rows=len(issues) + 1, cols=4)
            issue_table.style = "Table Grid"

            header = issue_table.rows[0]
            header.cells[0].text = "序号"
            header.cells[1].text = "问题描述"
            header.cells[2].text = "严重程度"
            header.cells[3].text = "状态"

            for i, issue in enumerate(issues):
                row = issue_table.rows[i + 1]
                row.cells[0].text = str(i + 1)
                row.cells[1].text = str(issue.get("description", ""))[:40]
                row.cells[2].text = str(issue.get("severity", "中"))
                row.cells[3].text = str(issue.get("status", "待处理"))
        else:
            doc.add_paragraph("本次测试未发现问题。")

        doc.save(output_path)

        return output_path

    def _generate_word_traceability_report(
        self, report_data: Dict[str, Any], output_path: str
    ) -> str:
        """生成Word溯源报告"""
        doc = Document()

        title = doc.add_heading(f"{report_data.get('project_name', '测试报告')} - 溯源报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("（内部审核文档）")
        doc.add_page_break()

        doc.add_heading("1. 数据来源信息", level=1)

        data_sources = report_data.get("data_sources", [])
        if data_sources:
            source_table = doc.add_table(rows=len(data_sources) + 1, cols=5)
            source_table.style = "Table Grid"

            header = source_table.rows[0]
            header.cells[0].text = "文件名"
            header.cells[1].text = "文件类型"
            header.cells[2].text = "大小"
            header.cells[3].text = "导入时间"
            header.cells[4].text = "数据质量"

            for i, source in enumerate(data_sources):
                row = source_table.rows[i + 1]
                row.cells[0].text = str(source.get("file_name", ""))[:25]
                row.cells[1].text = str(source.get("file_type", ""))
                row.cells[2].text = str(source.get("file_size", ""))
                row.cells[3].text = str(source.get("import_time", ""))
                row.cells[4].text = str(source.get("data_quality", "valid"))

        doc.add_heading("2. 信号溯源表", level=1)

        signal_mappings = report_data.get("signal_mappings", [])
        if signal_mappings:
            mapping_table = doc.add_table(rows=len(signal_mappings) + 1, cols=5)
            mapping_table.style = "Table Grid"

            header = mapping_table.rows[0]
            header.cells[0].text = "信号别名"
            header.cells[1].text = "原始信号"
            header.cells[2].text = "DBC信号"
            header.cells[3].text = "单位转换"
            header.cells[4].text = "数据来源"

            for i, mapping in enumerate(signal_mappings):
                row = mapping_table.rows[i + 1]
                unit_conv = mapping.get("unit_conversion", {})
                conv_str = ""
                if unit_conv:
                    conv_str = f"{unit_conv.get('from_unit', '')} -> {unit_conv.get('to_unit', '')}"

                row.cells[0].text = str(mapping.get("signal_alias", ""))[:15]
                row.cells[1].text = str(mapping.get("data_source_signal", ""))[:15]
                row.cells[2].text = str(mapping.get("dbc_signal", ""))[:15]
                row.cells[3].text = conv_str[:15]
                row.cells[4].text = str(mapping.get("source_file", ""))[:15]

        doc.add_heading("3. 详细分析结果", level=1)

        indicators = report_data.get("indicators", [])
        for i, ind in enumerate(indicators[:15]):
            doc.add_heading(f"3.{i + 1} {ind.get('indicator_id', 'Unknown')}", level=2)
            doc.add_paragraph(f"类型: {ind.get('indicator_type', 'N/A')}")
            doc.add_paragraph(f"状态: {ind.get('result_status', 'N/A')}")

            result_value = ind.get("result_value", {})
            if isinstance(result_value, dict):
                for key, value in result_value.items():
                    doc.add_paragraph(f"  - {key}: {value}")

            doc.add_paragraph(f"备注: {ind.get('notes', '')}")

        doc.add_heading("4. 原始数据引用", level=1)

        provenance = report_data.get("provenance", {})
        for key, value in provenance.items():
            doc.add_paragraph(f"  - {key}: {value}")

        doc.save(output_path)

        return output_path


report_engine = ReportEngine()
